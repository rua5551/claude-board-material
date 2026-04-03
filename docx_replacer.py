import os
import re
import sys
from docx.oxml.ns import qn
from docx import Document

def find_docx_files(folder_path):
    """查找文件夹中所有.docx文件（包括子文件夹）"""
    docx_files = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith('.docx'):
                docx_files.append(os.path.join(root, file))
                print(f"  发现文件: {os.path.join(root, file)}")
    return docx_files

def replace_in_docx(file_path, meeting_num, month, day):
    """替换docx文件内容并保留样式"""
    try:
        # --- 核心修复：确保日期是整数类型 ---
        month = int(month)
        day = int(day)
        # ----------------------------------
        doc = Document(file_path)
        modified = False

        # --- 准备目标文本 ---
        target_meeting = f'第{meeting_num}次'
        target_date = f'2026年{month}月{day}日'
        # 关键修复点：使用 rf 确保 \1 和 \2 能够正确代表正则分组
        target_doc_num = f'成金控董决〔2026〕{meeting_num}号'

        # --- 准备正则表达式 ---
        pattern_meeting = re.compile(r'第\d+次')
        pattern_date = re.compile(r'\d{4}年\d+月\d+日')
        # 关键修复点：匹配括号内的2026及后续数字，\s* 容错可能存在的空格
        pattern_doc_num = re.compile(r'成金控董决〔2026〕\d+号')

        def docx_safe_replace_logic(paragraph, search_re, replace_str):
            """合并Run逻辑以保留样式"""
            from docx.oxml.ns import qn
            
            inline = paragraph.runs
            text = "".join([run.text for run in inline])
            
            if search_re.search(text):
                # 进行正则替换
                new_text = search_re.sub(replace_str, text)
                
                if new_text != text:
                    # --- 改动 1：在清空前，先抓取原本的中文字体名称 ---
                    # rPr.rFonts.get(qn('w:eastAsia')) 可以直接获取 Word 底层记录的中文字体名
                    try:
                        original_chinese_font = inline[0]._element.rPr.rFonts.get(qn('w:eastAsia'))
                    except:
                        original_chinese_font = None # 如果抓不到，后面会用 font.name 代替

                    for run in inline:
                        run.text = ""
                        
                    if inline:
                        # --- 改动 2：写入新文本并设置西文（数字）字体 ---
                        inline[0].text = new_text
                        inline[0].font.name = 'Times New Roman'
                        
                        # --- 改动 3：将抓取到的原中文字体重新应用回去 ---
                        if original_chinese_font:
                            inline[0]._element.rPr.rFonts.set(qn('w:eastAsia'), original_chinese_font)
                        elif inline[0].font.name:
                            # 如果上面没抓到底层属性，尝试用通用属性补偿
                            inline[0]._element.rPr.rFonts.set(qn('w:eastAsia'), inline[0].font.name)
                            
                    return True
            return False

        def process_all_paragraphs(paragraphs):
            nonlocal modified
            num_p = len(paragraphs)
            for i in range(num_p):
                p = paragraphs[i]
                p_text = p.text # 获取当前行去掉空格的文本

                # 1. 替换会议次数 (放到判断外，确保无论如何都尝试替换)
                if pattern_meeting.search(p_text):
                    if docx_safe_replace_logic(p, pattern_meeting, target_meeting):
                        modified = True
                
                # 2. 替换日期 (增强版子规则判定)
                if pattern_date.search(p_text):
                    current_final_date = target_date
                    # 只有满足以下“落款条件”时，才执行 day-1
                    is_notice_date = False
                    
                    # 条件A：当前段落里就写着“联系人：”（兼容软回车 Shift+Enter）
                    if "联系人" in p_text:
                        is_notice_date = True
                    # 条件B：后续 3 行内出现了“联系人：”（兼容硬回车 Enter）
                    else:
                        for j in range(i + 1, min(i + 4, num_p)):
                            next_t = paragraphs[j].text.strip()
                            if not next_t: continue # 跳过空行
                            if "联系人" in next_t:
                                is_notice_date = True
                            break # 只要遇到第一个非空行，无论是否匹配都停止向下找
                    
                    # 如果判定为落款日期，执行减 1 逻辑
                    if is_notice_date:
                        if day > 1:
                            current_final_date = f"2026年{month}月{day-1}日"
                        else:
                            # 1号的情况无法自动减，保持原样并提示
                            print("-" * 50)
                            print(f"\n[!][!][!] 警告：文件 <{os.path.basename(file_path)}> 触发了联系人落款规则。")
                            print(f"    但你输入的日期是 1 号，无法自动减 1 天（总不能变 0 号）。")
                            print(f"    该文件的落款日期将保持为：2026年{month}月{day}日，请人工核对！\n")
                            print("-" * 50)

                    # 执行替换（不管是原日期还是减1后的日期，都执行一次替换以统一字体）
                    if docx_safe_replace_logic(p, pattern_date, current_final_date):
                        modified = True

                # 3. 替换特定文号
                if pattern_doc_num.search(p_text):
                    if docx_safe_replace_logic(p, pattern_doc_num, target_doc_num):
                        modified = True

        # 处理所有段落和表格
        process_all_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_all_paragraphs(cell.paragraphs)

        if modified:
            doc.save(file_path)
        return modified

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"处理文件 {file_path} 时出错: {e}")
        return False

def replace_in_filename(file_path, meeting_num, month, day):
    """替换文件名"""
    filename = os.path.basename(file_path)
    directory = os.path.dirname(file_path)

    new_filename = filename

    # 替换文件名中的"第n次"
    new_filename = re.sub(r'第(\d+)次', f'第{meeting_num}次', new_filename)

    # 替换文件名中的"2026年n月n日"
    new_filename = re.sub(r'2026年(\d+)月(\d+)日', f'2026年{month}月{day}日', new_filename)

    if new_filename != filename:
        new_path = os.path.join(directory, new_filename)
        if not os.path.exists(new_path):
            os.rename(file_path, new_path)
            return True
        else:
            print(f"目标文件名已存在: {new_path}")
            return False
    return False

def main():
    print("=" * 50)
    print("Word文档批量替换工具（限战发部用）")
    print("=" * 50)

    # 1. 输入文件夹路径
    folder_path = input("请输入文件夹路径（保证文件夹中为.docx文件）：").strip()
    folder_path = folder_path.strip('"').strip("'")

    if not os.path.isdir(folder_path):
        print("错误：指定的路径不是有效的文件夹！")
        input("按回车键退出...")
        return

    # 2. 输入第几次会议
    meeting_num = input("请输入第几次会议（数字）：").strip()
    if not meeting_num.isdigit():
        print("错误：请输入有效的数字！")
        input("按回车键退出...")
        return

    # 3. 输入月
    month = input("请输入月（数字）：").strip()
    if not month.isdigit() or not (1 <= int(month) <= 12):
        print("错误：请输入1-12之间的数字！")
        input("按回车键退出...")
        return

    # 4. 输入日
    day = input("请输入日（数字）：").strip()
    if not day.isdigit() or not (1 <= int(day) <= 31):
        print("错误：请输入1-31之间的数字！")
        input("回车键退出...")
        return

    print("\n开始处理...")
    print(f"文件夹: {folder_path}")
    print(f"会议次数: 第{meeting_num}次")
    print(f"日期: 2026年{month}月{day}日")
    print("-" * 50)

    # 查找所有docx文件
    docx_files = find_docx_files(folder_path)

    if not docx_files:
        print("未找到任何.docx文件！")
        input("关闭此程序退出...")
        return

    print(f"找到 {len(docx_files)} 个.docx文件\n")

    # 处理每个文件
    content_modified_count = 0
    filename_modified_count = 0

    for file_path in docx_files:
        filename = os.path.basename(file_path)
        print(f"处理: {filename}")

        # 替换内容
        if replace_in_docx(file_path, meeting_num, month, day):
            content_modified_count += 1
            print(f"  - 内容已替换")

        # 替换文件名
        if replace_in_filename(file_path, meeting_num, month, day):
            filename_modified_count += 1
            print(f"  - 文件名已修改")

    print("-" * 50)
    print(f"- 内容已替换: {content_modified_count} 个文件")
    print(f"- 文件名已修改: {filename_modified_count} 个文件")
    print("\n请继续修改：\n1.《会议通知》中会议时间“（星期X）”\n2.《会议通知》《会议决议》中具体议案部分")
    print("\n-关闭此程序退出...")
    input()

if __name__ == "__main__":
    main()